import fitz  # PyMuPDF
import docx
import io

# def read_pdf(blob_data):
#     with open("temp.pdf", "wb") as temp_file:
#         temp_file.write(blob_data)
#     document = fitz.open("temp.pdf")
#     text = ""
#     for page_num in range(len(document)):
#         page = document.load_page(page_num)
#         text += page.get_text()
#     return text

# def read_docx(blob_data):
#     with open("temp.docx", "wb") as temp_file:
#         temp_file.write(blob_data)
#     doc = docx.Document("temp.docx")
#     text = ""
#     for para in doc.paragraphs:
#         text += para.text + "\n"
#     return text

def read_pdf(blob_data):
    with io.BytesIO(blob_data) as temp_file:
        document = fitz.open("pdf", temp_file)
        text = ""
        for page_num in range(len(document)):
            page = document.load_page(page_num)
            text += page.get_text()
    return text

def read_docx(blob_data):
    with io.BytesIO(blob_data) as temp_file:
        doc = docx.Document(temp_file)
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
    return text